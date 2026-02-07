import streamlit as st
from datetime import datetime, date
from io import BytesIO

from PIL import Image

# PDF (ReportLab - visual)
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

# Word (python-docx)
from docx import Document
from docx.shared import Inches

# Excel (openpyxl)
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.datavalidation import DataValidation


st.set_page_config(
    page_title="Control Edificio Pro (Streamlit)",
    page_icon="🛡️",
    layout="wide",
)

# ---------------------------
# Master data (defaults)
# ---------------------------
CATEGORIES = ["Críticos", "Accesos", "Higiene", "Comunes", "Infra"]

DEFAULT_INSTALLATIONS = [
    {"Tipo": "Críticos", "Instalación": "Sala de Bombas", "Tarea": "Presión y alternancia"},
    {"Tipo": "Críticos", "Instalación": "Sala de Calderas", "Tarea": "Temperatura y fugas"},
    {"Tipo": "Críticos", "Instalación": "Generador", "Tarea": "Nivel petróleo y batería"},
    {"Tipo": "Críticos", "Instalación": "PEAS (Presurización)", "Tarea": "Prueba de ventilador"},
    {"Tipo": "Críticos", "Instalación": "Ascensores (2)", "Tarea": "Nivelación y limpieza rieles"},
    {"Tipo": "Accesos", "Instalación": "Portones (2)", "Tarea": "Sensores y velocidad"},
    {"Tipo": "Accesos", "Instalación": "Control Biométrico", "Tarea": "Lectores huella/tarjeta"},
    {"Tipo": "Higiene", "Instalación": "Sala de Basura", "Tarea": "Desinfección y contenedores"},
    {"Tipo": "Higiene", "Instalación": "Ductos (20 pisos)", "Tarea": "Cierre de escotillas"},
    {"Tipo": "Comunes", "Instalación": "Piscina", "Tarea": "Parámetros Cl/pH"},
    {"Tipo": "Comunes", "Instalación": "Quincho / Eventos", "Tarea": "Mobiliario e higiene"},
    {"Tipo": "Comunes", "Instalación": "Gym / Sauna", "Tarea": "Máquinas y tableros"},
    {"Tipo": "Infra", "Instalación": "Pasillos (1-20)", "Tarea": "Luces de emergencia"},
    {"Tipo": "Infra", "Instalación": "Subterráneo", "Tarea": "Filtraciones y limpieza"},
    {"Tipo": "Infra", "Instalación": "Jardines", "Tarea": "Riego programado"},
]


# ---------------------------
# Helpers (State)
# ---------------------------
def build_checklist_items_from_master(master_rows):
    """
    master_rows: list of dicts with keys Tipo, Instalación, Tarea (optional)
    Builds session checklist item dicts (id, cat, name, task, status, note, photo)
    """
    items = []
    next_id = 1
    for r in master_rows:
        cat = (r.get("Tipo") or "").strip()
        name = (r.get("Instalación") or "").strip()
        task = (r.get("Tarea") or "").strip() or "—"

        if not cat or not name:
            continue

        # Normaliza cat a las categorías permitidas si viene con variantes
        if cat not in CATEGORIES:
            # Si viene "Espacio Común" etc., intenta mapear a Comunes
            mapped = map_tipo_to_category(cat)
            cat = mapped

        items.append({
            "id": next_id,
            "cat": cat,
            "name": name,
            "task": task,
            "status": "pending",
            "note": "",
            "photo": None,
        })
        next_id += 1

    return items


def map_tipo_to_category(tipo: str) -> str:
    """
    Mapea 'Tipo' libre a una categoría soportada.
    """
    t = (tipo or "").strip().lower()
    if "crit" in t:
        return "Críticos"
    if "infra" in t:
        return "Infra"
    if "comun" in t or "común" in t or "espacio" in t:
        return "Comunes"
    if "hig" in t or "aseo" in t or "basura" in t:
        return "Higiene"
    if "acces" in t or "port" in t:
        return "Accesos"
    # fallback
    return "Comunes"


def init_state():
    if "community_name" not in st.session_state:
        st.session_state["community_name"] = "Comunidad (sin nombre)"

    if "report_date" not in st.session_state:
        st.session_state["report_date"] = date.today()

    # checklist_items: se construye desde defaults al inicio
    if "checklist_items" not in st.session_state:
        st.session_state["checklist_items"] = build_checklist_items_from_master(DEFAULT_INSTALLATIONS)

    if "incidences" not in st.session_state:
        st.session_state["incidences"] = []  # {id, employee, detail, ts}

    if "needs" not in st.session_state:
        st.session_state["needs"] = ""


# ---------------------------
# Helpers (UI/Stats/Text)
# ---------------------------
def status_badge(status: str) -> str:
    if status == "ok":
        return "🟢 OK"
    if status == "fail":
        return "🔴 FALLA"
    return "⚪ PEND."


def status_color(status: str) -> str:
    if status == "ok":
        return "#16a34a"
    if status == "fail":
        return "#dc2626"
    return "#64748b"


def get_stats():
    items = st.session_state["checklist_items"]
    ok = sum(1 for i in items if i["status"] == "ok")
    fail = sum(1 for i in items if i["status"] == "fail")
    pending = sum(1 for i in items if i["status"] == "pending")
    total = len(items)
    return ok, fail, pending, total


def build_report_text():
    rd = st.session_state["report_date"]
    community = st.session_state["community_name"]
    ok, fail, pending, total = get_stats()

    lines = []
    lines.append("--- INFORME DE GESTIÓN / CONTROL DE INSTALACIONES ---")
    lines.append(f"COMUNIDAD: {community}")
    lines.append(f"FECHA: {rd.isoformat()}")
    lines.append("")
    lines.append(f"RESUMEN: OK={ok} | FALLAS={fail} | PENDIENTES={pending} | TOTAL={total}")
    lines.append("")
    lines.append("1) CHECKLIST TÉCNICO")
    lines.append("---------------------------------")

    for cat in CATEGORIES:
        lines.append(f"\n[{cat}]")
        for it in [x for x in st.session_state["checklist_items"] if x["cat"] == cat]:
            note = (it.get("note") or "").strip()
            note_txt = note if note else "Sin novedades."
            lines.append(f"- {status_badge(it['status'])} {it['name']} ({it['task']}): {note_txt}")

    lines.append("\n2) REQUERIMIENTOS / COMPRAS")
    lines.append("---------------------------------")
    needs = (st.session_state["needs"] or "").strip() or "Sin requerimientos reportados."
    lines.append(needs)

    lines.append("\n3) INCIDENCIAS RR.HH.")
    lines.append("---------------------------------")
    incidences = st.session_state["incidences"]
    if not incidences:
        lines.append("Sin incidencias registradas.")
    else:
        for inc in sorted(incidences, key=lambda x: x["ts"], reverse=True):
            ts = inc["ts"].strftime("%Y-%m-%d %H:%M")
            lines.append(f"- {ts} | {inc['employee']}: {inc['detail']}")

    return "\n".join(lines)


# ---------------------------
# Excel (Master data template)
# ---------------------------
def export_master_template_bytes() -> bytes:
    """
    Crea plantilla XLSX para Datos Maestros:
    Columnas: Tipo, Instalación, Tarea (opcional)
    Tipo tiene validación (dropdown) con categorías soportadas.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "DatosMaestros"

    headers = ["Tipo", "Instalación", "Tarea"]
    ws.append(headers)

    # precarga con defaults (si el usuario quiere usarlos)
    for r in DEFAULT_INSTALLATIONS:
        ws.append([r.get("Tipo", ""), r.get("Instalación", ""), r.get("Tarea", "")])

    # Ajustes simples
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 32
    ws.column_dimensions["C"].width = 34

    # Validación de datos (dropdown) en Tipo
    allowed = ",".join(CATEGORIES)
    dv = DataValidation(type="list", formula1=f'"{allowed}"', allow_blank=False)
    ws.add_data_validation(dv)

    # Aplica validación a un rango “amplio” (por si agregan filas)
    dv.add("A2:A500")

    # Guarda a bytes
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.getvalue()


def import_master_from_xlsx(uploaded_file_bytes: bytes):
    """
    Lee XLSX y reemplaza checklist_items por lo que venga en la hoja DatosMaestros (o la primera hoja).
    Espera columnas: Tipo, Instalación, (Tarea opcional).
    """
    bio = BytesIO(uploaded_file_bytes)
    wb = load_workbook(bio, data_only=True)

    if "DatosMaestros" in wb.sheetnames:
        ws = wb["DatosMaestros"]
    else:
        ws = wb[wb.sheetnames[0]]

    # leer headers
    rows = list(ws.iter_rows(values_only=True))
    if not rows or len(rows) < 2:
        raise ValueError("La plantilla no contiene datos.")

    header = [str(x).strip() if x is not None else "" for x in rows[0]]
    # indices
    def idx(col):
        try:
            return header.index(col)
        except ValueError:
            return None

    i_tipo = idx("Tipo")
    i_inst = idx("Instalación")
    i_task = idx("Tarea")  # opcional

    if i_tipo is None or i_inst is None:
        raise ValueError("La plantilla debe incluir columnas 'Tipo' y 'Instalación'.")

    master = []
    for r in rows[1:]:
        if r is None:
            continue
        tipo = (r[i_tipo] if i_tipo < len(r) else "") or ""
        inst = (r[i_inst] if i_inst < len(r) else "") or ""
        task = (r[i_task] if (i_task is not None and i_task < len(r)) else "") or ""
        tipo = str(tipo).strip()
        inst = str(inst).strip()
        task = str(task).strip()

        if not tipo or not inst:
            continue

        master.append({"Tipo": tipo, "Instalación": inst, "Tarea": task})

    if not master:
        raise ValueError("No se encontraron registros válidos (Tipo + Instalación).")

    st.session_state["checklist_items"] = build_checklist_items_from_master(master)


# ---------------------------
# PDF (Visual 3-column table + red soft background on FAIL + summary cards)
# ---------------------------
def _status_tag_html(status: str) -> str:
    if status == "ok":
        return '<font color="#16a34a"><b>OK</b></font>'
    if status == "fail":
        return '<font color="#dc2626"><b>FALLA</b></font>'
    return '<font color="#64748b"><b>PEND.</b></font>'


def _make_rl_image(photo_bytes: bytes, max_w: float, max_h: float, small_style):
    if not photo_bytes:
        return Paragraph("<i>Sin foto</i>", small_style)

    try:
        img = Image.open(BytesIO(photo_bytes)).convert("RGB")
    except Exception:
        return Paragraph("<i>Foto inválida</i>", small_style)

    iw, ih = img.size
    scale = min(max_w / iw, max_h / ih)
    w, h = iw * scale, ih * scale

    buf = BytesIO()
    img.save(buf, format="JPEG", quality=85)
    buf.seek(0)

    from reportlab.platypus import Image as RLImage
    return RLImage(buf, width=w, height=h)


def generate_pdf_bytes_visual() -> bytes:
    styles = getSampleStyleSheet()

    normal = ParagraphStyle(
        "normal",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#0f172a"),
    )
    small = ParagraphStyle(
        "small",
        parent=normal,
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#334155"),
    )
    title_style = ParagraphStyle(
        "title_style",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=18,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=8,
    )
    cat_style = ParagraphStyle(
        "cat_style",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12,
        textColor=colors.HexColor("#4338ca"),
        spaceBefore=10,
        spaceAfter=6,
    )

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )

    elements = []
    rd = st.session_state["report_date"]
    community = st.session_state["community_name"]
    ok, fail, pending, total = get_stats()

    # Title
    elements.append(Paragraph("Control Edificio Pro – Informe Visual", title_style))
    elements.append(Paragraph(f"<b>Comunidad:</b> {community}", normal))
    elements.append(Paragraph(f"<b>Fecha:</b> {rd.isoformat()}", normal))
    elements.append(Spacer(1, 8))

    # Summary cards
    summary_data = [
        [
            Paragraph("<b>Sistemas OK</b>", small),
            Paragraph("<b>Fallas</b>", small),
            Paragraph("<b>Pendientes</b>", small),
            Paragraph("<b>Total</b>", small),
        ],
        [
            Paragraph(f"<font color='#16a34a'><b>{ok}</b></font>", normal),
            Paragraph(f"<font color='#dc2626'><b>{fail}</b></font>", normal),
            Paragraph(f"<font color='#64748b'><b>{pending}</b></font>", normal),
            Paragraph(f"<b>{total}</b>", normal),
        ],
    ]
    summary_table = Table(summary_data, colWidths=[4.2 * cm, 4.2 * cm, 4.2 * cm, 4.2 * cm])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2ff")),
        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#ffffff")),
        ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#e2e8f0")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    elements.append(summary_table)
    elements.append(Spacer(1, 12))

    # Table layout
    col1_w = 6.2 * cm
    col2_w = 6.2 * cm
    col3_w = 5.6 * cm
    photo_max_w = col3_w - 0.3 * cm
    photo_max_h = 3.6 * cm

    items = st.session_state["checklist_items"]

    for cat in CATEGORIES:
        cat_items = [x for x in items if x["cat"] == cat]
        if not cat_items:
            continue

        elements.append(Paragraph(cat, cat_style))

        data = [
            [
                Paragraph("<b>Instalación</b>", normal),
                Paragraph("<b>Estado / Observación</b>", normal),
                Paragraph("<b>Registro visual</b>", normal),
            ]
        ]

        fail_row_indices = []

        for it in cat_items:
            inst = Paragraph(
                f"<b>{it['name']}</b><br/><font color='#64748b'>{it['task']}</font>",
                normal,
            )
            note = (it.get("note") or "").strip()
            note_txt = note if note else "Sin novedades."

            mid = Paragraph(
                f"{_status_tag_html(it['status'])}<br/>{note_txt}",
                normal,
            )

            img_cell = _make_rl_image(it.get("photo"), photo_max_w, photo_max_h, small)
            data.append([inst, mid, img_cell])

            if it["status"] == "fail":
                fail_row_indices.append(len(data) - 1)

        table = Table(data, colWidths=[col1_w, col2_w, col3_w])

        style_cmds = [
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2ff")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("VALIGN", (0, 1), (-1, -1), "TOP"),
        ]

        soft_red = colors.HexColor("#fee2e2")
        for r in fail_row_indices:
            style_cmds.append(("BACKGROUND", (0, r), (-1, r), soft_red))

        table.setStyle(TableStyle(style_cmds))

        elements.append(table)
        elements.append(Spacer(1, 10))

    # Footer sections
    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Requerimientos / Compras", cat_style))
    needs = (st.session_state["needs"] or "").strip() or "Sin requerimientos reportados."
    elements.append(Paragraph(needs, normal))

    elements.append(Spacer(1, 6))
    elements.append(Paragraph("Incidencias RR.HH.", cat_style))
    incidences = st.session_state["incidences"]
    if not incidences:
        elements.append(Paragraph("Sin incidencias registradas.", normal))
    else:
        for inc in sorted(incidences, key=lambda x: x["ts"], reverse=True):
            ts = inc["ts"].strftime("%Y-%m-%d %H:%M")
            elements.append(Paragraph(f"- <b>{ts}</b> | <b>{inc['employee']}</b>: {inc['detail']}", normal))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------
# Word (texto + anexo fotos, mantiene tu versión actual)
# ---------------------------
def generate_docx_bytes(report_text: str) -> bytes:
    doc = Document()
    doc.add_heading("Informe de Gestión / Control de Instalaciones", level=1)

    for line in report_text.split("\n"):
        doc.add_paragraph(line)

    items = st.session_state["checklist_items"]
    photos = [it for it in items if it.get("photo")]

    if photos:
        doc.add_page_break()
        doc.add_heading("Anexo: Fotos", level=2)

        for it in photos:
            doc.add_paragraph(f"#{it['id']} - {it['cat']} - {it['name']} ({it['task']})")
            note = (it.get("note") or "").strip()
            doc.add_paragraph(f"Obs: {note}" if note else "Obs: (sin observaciones)")

            img = Image.open(BytesIO(it["photo"])).convert("RGB")
            img_buf = BytesIO()
            img.save(img_buf, format="PNG")
            img_buf.seek(0)

            doc.add_picture(img_buf, width=Inches(5.8))
            doc.add_paragraph("")

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


# ---------------------------
# UI
# ---------------------------
init_state()
ok, fail, pending, total = get_stats()

# Header
st.markdown(
    f"""
    <div style="padding:18px 18px 10px 18px; border-radius:16px; background: linear-gradient(90deg, #4338ca, #4f46e5); color:white;">
      <div style="display:flex; justify-content:space-between; align-items:center; gap:16px; flex-wrap:wrap;">
        <div>
          <div style="font-size:24px; font-weight:800;">🛡️ Control Edificio Pro</div>
          <div style="opacity:0.85;">Mayordomía y Gestión de Operaciones (demo Streamlit)</div>
        </div>
        <div style="display:flex; gap:16px;">
          <div style="background: rgba(0,0,0,0.20); padding:10px 14px; border-radius:12px; text-align:center; min-width:120px;">
            <div style="font-size:11px; letter-spacing:0.08em; opacity:0.8;">SISTEMAS OK</div>
            <div style="font-size:24px; font-weight:800; color:#4ade80;">{ok}</div>
          </div>
          <div style="background: rgba(0,0,0,0.20); padding:10px 14px; border-radius:12px; text-align:center; min-width:120px;">
            <div style="font-size:11px; letter-spacing:0.08em; opacity:0.8;">FALLAS</div>
            <div style="font-size:24px; font-weight:800; color:#f87171;">{fail}</div>
          </div>
        </div>
      </div>
    </div>
    """,
    unsafe_allow_html=True
)
st.write("")

# Tabs
tab_checklist, tab_rrhh, tab_report, tab_master = st.tabs(
    ["✅ Levantamiento Técnico", "👥 RR.HH. (Incidencias)", "🧾 Generar Informe", "⚙️ Datos Maestros (Instalaciones)"]
)

# ---------------------------
# Checklist
# ---------------------------
with tab_checklist:
    c1, c2 = st.columns([2, 1])
    with c1:
        st.subheader("Checklist Técnico (por áreas)")
        st.caption("Marca OK / FALLA / PENDIENTE, agrega observaciones y (opcional) una foto por ítem.")
    with c2:
        st.session_state["report_date"] = st.date_input("Fecha del informe", value=st.session_state["report_date"])

    st.session_state["community_name"] = st.text_input(
        "Nombre de la comunidad",
        value=st.session_state["community_name"],
        placeholder="Ej: Edificio Los Castaños 123",
    )

    for cat in CATEGORIES:
        st.markdown(f"### {cat}")
        for it in [x for x in st.session_state["checklist_items"] if x["cat"] == cat]:
            box = st.container(border=True)
            with box:
                left, mid, right = st.columns([2.2, 2.2, 1.6], gap="medium")

                with left:
                    st.markdown(f"**{it['name']}**")
                    st.caption(it["task"])

                with mid:
                    status = st.radio(
                        "Estado",
                        options=["pending", "ok", "fail"],
                        format_func=lambda v: {"pending": "Pendiente", "ok": "OK", "fail": "Falla"}[v],
                        index=["pending", "ok", "fail"].index(it["status"]),
                        horizontal=True,
                        key=f"status_{it['id']}",
                        label_visibility="collapsed",
                    )
                    it["status"] = status

                    it["note"] = st.text_input(
                        "Observación",
                        value=it["note"],
                        placeholder="Escribe una observación breve…",
                        key=f"note_{it['id']}",
                        label_visibility="collapsed",
                    )

                with right:
                    st.markdown(
                        f"""
                        <div style="padding:10px 12px; border-radius:12px; border:1px solid #e2e8f0;">
                          <div style="font-weight:800; color:{status_color(it['status'])}; font-size:16px;">
                            {status_badge(it['status'])}
                          </div>
                          <div style="opacity:0.7; font-size:12px;">Ítem #{it['id']}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                    uploaded = st.file_uploader(
                        "Foto (opcional)",
                        type=["png", "jpg", "jpeg"],
                        key=f"photo_{it['id']}",
                        label_visibility="collapsed",
                    )
                    if uploaded is not None:
                        it["photo"] = uploaded.getvalue()
                        st.image(it["photo"], caption="Foto adjunta", use_container_width=True)

        st.divider()

    colA, colB, colC = st.columns([1, 1, 2])
    with colA:
        if st.button("🔄 Marcar todo como Pendiente"):
            for it in st.session_state["checklist_items"]:
                it["status"] = "pending"
            st.rerun()

    with colB:
        if st.button("✅ Marcar todo como OK"):
            for it in st.session_state["checklist_items"]:
                it["status"] = "ok"
            st.rerun()

    with colC:
        st.info("Tip: sin base de datos, el contenido se mantiene mientras no cierres/recargues la pestaña. Para persistencia real, se integra Firebase/Sheets.")


# ---------------------------
# RR.HH - Incidences
# ---------------------------
with tab_rrhh:
    st.subheader("Gestión RR.HH. – Incidencias manuales")
    st.caption("Registra incidencias por empleado (puedes agregar múltiples incidencias para el mismo nombre).")

    with st.form("add_incidence", clear_on_submit=True):
        employee = st.text_input("Nombre del empleado", placeholder="Ej: Juan Pérez")
        detail = st.text_area("Detalle de la incidencia", placeholder="Ej: Atraso 20 min. / Falta de EPP / Observación…", height=120)
        submitted = st.form_submit_button("➕ Ingresar nueva incidencia")

        if submitted:
            if not employee.strip() or not detail.strip():
                st.error("Por favor completa nombre del empleado y detalle.")
            else:
                incidences = st.session_state["incidences"]
                new_id = (max([i["id"] for i in incidences]) + 1) if incidences else 1
                st.session_state["incidences"].append(
                    {"id": new_id, "employee": employee.strip(), "detail": detail.strip(), "ts": datetime.now()}
                )
                st.success("Incidencia registrada.")

    st.write("")
    st.markdown("#### Incidencias registradas")

    if not st.session_state["incidences"]:
        st.warning("Aún no hay incidencias.")
    else:
        for inc in sorted(st.session_state["incidences"], key=lambda x: x["ts"], reverse=True):
            with st.container(border=True):
                c1, c2, c3 = st.columns([2, 6, 1.2])
                with c1:
                    st.markdown(f"**{inc['employee']}**")
                    st.caption(inc["ts"].strftime("%Y-%m-%d %H:%M"))
                with c2:
                    st.write(inc["detail"])
                with c3:
                    if st.button("🗑️", key=f"del_inc_{inc['id']}", help="Eliminar incidencia"):
                        st.session_state["incidences"] = [x for x in st.session_state["incidences"] if x["id"] != inc["id"]]
                        st.rerun()


# ---------------------------
# Report
# ---------------------------
with tab_report:
    st.subheader("Generador de Informe (descarga PDF o Word con fotos)")
    st.caption("El PDF se exporta con estructura visual (3 columnas por área). El Word se exporta en texto + anexo de fotos (por ahora).")

    st.session_state["needs"] = st.text_area(
        "Requerimientos y compras (texto libre)",
        value=st.session_state["needs"],
        height=140,
        placeholder="Ej: Solicitar mantención ascensores / compra de luminarias / repuestos bomba…",
    )

    report_text = build_report_text()

    st.markdown("#### Vista previa (texto)")
    st.code(report_text, language="text")

    fmt = st.radio("Formato de descarga", options=["PDF (Visual)", "Word (DOCX)"], horizontal=True)

    file_base = f"informe_{st.session_state['community_name'].strip().replace(' ', '_')}_{st.session_state['report_date'].isoformat()}"

    col1, col2 = st.columns([1, 2])
    with col1:
        if fmt.startswith("PDF"):
            pdf_bytes = generate_pdf_bytes_visual()
            st.download_button(
                "⬇️ Descargar PDF (Visual, con fotos)",
                data=pdf_bytes,
                file_name=f"{file_base}.pdf",
                mime="application/pdf",
            )
        else:
            docx_bytes = generate_docx_bytes(report_text)
            st.download_button(
                "⬇️ Descargar Word (DOCX) (con fotos)",
                data=docx_bytes,
                file_name=f"{file_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    with col2:
        st.info("Tip: si subes muchas fotos grandes, el PDF/DOCX puede quedar pesado. Si quieres, agregamos compresión automática.")


# ---------------------------
# Master Data (Instalaciones)
# ---------------------------
with tab_master:
    st.subheader("Datos Maestros de Instalaciones")
    st.caption("Puedes descargar una plantilla XLSX, editarla y volver a cargarla para personalizar el checklist. También puedes agregar/eliminar instalaciones aquí mismo.")

    # Descargar plantilla
    template_bytes = export_master_template_bytes()
    st.download_button(
        "⬇️ Descargar plantilla XLSX (Datos Maestros)",
        data=template_bytes,
        file_name="plantilla_datos_maestros_instalaciones.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    st.write("")
    # Subir plantilla
    uploaded_xlsx = st.file_uploader(
        "Cargar plantilla XLSX (reemplaza las instalaciones actuales)",
        type=["xlsx"],
        help="Debe contener columnas: Tipo, Instalación (Tarea opcional).",
    )

    if uploaded_xlsx is not None:
        try:
            import_master_from_xlsx(uploaded_xlsx.getvalue())
            st.success("Datos maestros cargados. Se actualizó el checklist.")
            st.rerun()
        except Exception as e:
            st.error(f"No se pudo cargar la plantilla: {e}")

    st.divider()

    # Agregar instalación manual
    st.markdown("### ➕ Agregar instalación (manual)")
    with st.form("add_installation"):
        colA, colB, colC = st.columns([1.3, 2.2, 2.2])
        with colA:
            tipo = st.selectbox("Tipo", options=CATEGORIES)
        with colB:
            instalacion = st.text_input("Instalación", placeholder="Ej: Sala de Tableros")
        with colC:
            tarea = st.text_input("Tarea (opcional)", placeholder="Ej: Revisión térmica / limpieza / fugas")

        add = st.form_submit_button("Agregar")
        if add:
            if not instalacion.strip():
                st.error("Debes indicar el nombre de la instalación.")
            else:
                items = st.session_state["checklist_items"]
                new_id = max([x["id"] for x in items], default=0) + 1
                items.append({
                    "id": new_id,
                    "cat": tipo,
                    "name": instalacion.strip(),
                    "task": tarea.strip() or "—",
                    "status": "pending",
                    "note": "",
                    "photo": None,
                })
                st.success("Instalación agregada.")
                st.rerun()

    st.divider()

    # Eliminar instalaciones
    st.markdown("### 🗑️ Quitar instalaciones")
    items = st.session_state["checklist_items"]
    options = [f"#{x['id']} | {x['cat']} | {x['name']}" for x in items]
    to_remove = st.multiselect("Selecciona instalaciones a eliminar", options=options)

    if st.button("Eliminar seleccionadas"):
        if not to_remove:
            st.warning("No seleccionaste ninguna.")
        else:
            ids = set()
            for s in to_remove:
                # formato: "#id | ..."
                try:
                    part = s.split("|")[0].strip()
                    rid = int(part.replace("#", "").strip())
                    ids.add(rid)
                except:
                    pass

            st.session_state["checklist_items"] = [x for x in st.session_state["checklist_items"] if x["id"] not in ids]
            # Reasignar IDs (limpio y ordenado)
            rebuilt = []
            nid = 1
            for x in st.session_state["checklist_items"]:
                x["id"] = nid
                rebuilt.append(x)
                nid += 1
            st.session_state["checklist_items"] = rebuilt

            st.success("Instalaciones eliminadas.")
            st.rerun()

    st.divider()

    # Reset defaults
    st.markdown("### 🔁 Restaurar instalaciones por defecto")
    if st.button("Restaurar checklist por defecto (precargado)"):
        st.session_state["checklist_items"] = build_checklist_items_from_master(DEFAULT_INSTALLATIONS)
        st.success("Restaurado.")
        st.rerun()


st.markdown(
    "<div style='opacity:0.6; font-size:12px; margin-top:18px;'>Plataforma de Control Interno v1.0 (demo) • Streamlit</div>",
    unsafe_allow_html=True
)

